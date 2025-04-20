import os
import re
from tqdm import tqdm
from docx import Document as DocxDocument
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain.embeddings import HuggingFaceEmbeddings

# Configuration
BASE_VECTOR_STORE = "vector_store"
DATA_FOLDER = "data"
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
CHUNK_SIZE = 500
MIN_CHUNK_LENGTH = 100
COMBINED_INDEX_NAME = "db_faiss"

# ✅ Clean text
def clean_text(text):
    if not text:
        return ""
    text = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', text)
    text = re.sub(r'(?<=\w)-\s+(?=\w)', '', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[•▪�]+', ' ', text)
    text = re.sub(r'(?i)\b(?:page|pase)\s*\d+\b', '', text)
    text = re.sub(r'_{2,}', '', text)
    return text.strip()

# ✅ Detect item number and resolution
def detect_item_and_resolution(text):
    item_match = re.search(r'Item No\.\s*\d+', text, re.IGNORECASE)
    resolution_match = re.search(r'Resolution with respect to[:\-]?\s*(.+)', text, re.IGNORECASE)
    return (
        item_match.group(0) if item_match else None,
        resolution_match.group(1).strip() if resolution_match else None
    )

# ✅ Process DOCX and chunk per item
def process_docx_by_items(docx_path):
    doc = DocxDocument(docx_path)
    file_name = os.path.basename(docx_path)

    chunks = []
    current_item_no = None
    current_resolution = None
    current_text = []

    def store_chunk():
        nonlocal current_text
        if current_text:
            text_block = "\n".join(current_text).strip()
            if len(text_block) >= MIN_CHUNK_LENGTH:
                metadata = {
                    "source": file_name,
                    "item_no": current_item_no,
                    "resolution": current_resolution
                }
                chunks.append(Document(page_content=text_block, metadata=metadata))
            current_text = []

    # Process paragraphs
    for para in doc.paragraphs:
        cleaned = clean_text(para.text)
        if not cleaned:
            continue

        item, resolution = detect_item_and_resolution(cleaned)

        if item:
            store_chunk()
            current_item_no = item
            current_resolution = resolution if resolution else current_resolution
        elif resolution:
            current_resolution = resolution

        current_text.append(cleaned)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(clean_text(cell.text) for cell in row.cells if clean_text(cell.text))
            if row_text:
                current_text.append(row_text)

    # Final flush
    store_chunk()

    return chunks

# ✅ Store vectors
def store_per_doc_embeddings():
    embedder = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
    files = [f for f in os.listdir(DATA_FOLDER) if f.lower().endswith('.docx')]

    if not files:
        print("No DOCX files found.")
        return

    all_chunks = []

    for file in tqdm(files, desc="Processing DOCX Files"):
        file_path = os.path.join(DATA_FOLDER, file)
        chunks = process_docx_by_items(file_path)

        if not chunks:
            print(f"⚠️ No valid chunks for {file}")
            continue

        all_chunks.extend(chunks)

        try:
            vector_store = FAISS.from_documents(chunks, embedder)
        except Exception as e:
            print(f"❌ Failed to embed/index {file}: {e}")
            continue

        doc_vector_dir = os.path.join(BASE_VECTOR_STORE, os.path.splitext(file)[0])
        os.makedirs(doc_vector_dir, exist_ok=True)
        vector_store.save_local(doc_vector_dir)
        print(f"✅ Stored FAISS index for {file} in {doc_vector_dir}")

    # ✅ Store combined db_faiss
    if all_chunks:
        try:
            print("\n💾 Saving combined vector store: db_faiss...")
            combined_store = FAISS.from_documents(all_chunks, embedder)
            combined_path = os.path.join(BASE_VECTOR_STORE, COMBINED_INDEX_NAME)
            os.makedirs(combined_path, exist_ok=True)
            combined_store.save_local(combined_path)
            print(f"✅ Stored combined FAISS index at: {combined_path}")
        except Exception as e:
            print(f"❌ Failed to store combined FAISS index: {e}")

if __name__ == "__main__":
    store_per_doc_embeddings()
